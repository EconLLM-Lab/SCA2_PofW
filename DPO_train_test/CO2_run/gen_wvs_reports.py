#!/usr/bin/env python3
"""Generate (1) Ksennia-style per-adapter distribution-grid docx and
(2) an executive report docx with metrics tables, key graphs, and a
comparison against Ksennia's qualitative results.
Runs on the Colab VM (data on Drive). Outputs to /content/outputs/.
"""
import pathlib, math
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

R = pathlib.Path("/content/drive/MyDrive/DPO_CO2/eval_results_wvs_wave7")
QMAP = pathlib.Path("/content/wvs_eval/question_map_wvs_edited.csv")
OUT = pathlib.Path("/content/outputs")
OUT.mkdir(parents=True, exist_ok=True)

CCS = ["IND", "IDN", "NGA", "EGY", "TUR", "NLD", "BRA", "GRC"]
DIMS = ["Unspecified", "altruism", "negrecip", "patience", "posrecip", "risktaking", "trust"]
BLUE, ORANGE = "#1f77b4", "#ff7f0e"
GREEN, RED = "#2ca02c", "#d62728"

msl = pd.read_csv(R / "model_option_probabilities.csv")
pop = pd.read_csv(R / "population_response_distributions.csv")
qm = pd.read_csv(R / "survey_question_metrics_all_models.csv")
boot = pd.read_csv(R / "survey_metric_bootstrap_summary.csv")
ece = pd.read_csv(R / "survey_probability_calibration_ece.csv")
imp = pd.read_csv(R / "survey_adapter_improvement_vs_base_by_gps_dimension.csv")
qmap = pd.read_csv(QMAP).set_index("Question")

RR_ACC = {"IND": 0.947, "IDN": 0.886, "NGA": 0.977, "EGY": 0.886,
          "TUR": 0.955, "NLD": 0.932, "BRA": 0.939, "GRC": 0.947}

def qdim(q):
    if q in qmap.index:
        d = qmap.loc[q, "gps_dimension"]
        return d if pd.notna(d) else "Unspecified"
    return "Unspecified"

def qtext(q):
    if q in qmap.index:
        t = str(qmap.loc[q, "QuestionText"])
        return t[:46] + ("…" if len(t) > 46 else "")
    return ""

# ---------------------------------------------------------------- matched table
matched_rows = []
for c in CCS:
    b = qm[(qm.model == "base") & (qm.eval_country == c)].set_index("question_id")["tv_distance"]
    a = qm[(qm.model == f"{c}_adapter") & (qm.eval_country == c)].set_index("question_id")["tv_distance"]
    d = (b - a).dropna()
    matched_rows.append({
        "country": c, "n": len(d),
        "base_tvd": b.mean(), "adapter_tvd": a.mean(), "delta": d.mean(),
        "frac_improved": (d > 0).mean(),
        "base_jsd": qm[(qm.model == "base") & (qm.eval_country == c)]["js_divergence"].mean(),
        "adapter_jsd": qm[(qm.model == f"{c}_adapter") & (qm.eval_country == c)]["js_divergence"].mean(),
    })
mt = pd.DataFrame(matched_rows).set_index("country")

# ---------------------------------------------------------------- Ksennia-style grids
def adapter_grid(country, path):
    c = country
    popc = pop[pop.eval_country == c]
    if "eval_country" in msl.columns:
        mod = msl[(msl.model == f"{c}_adapter") & (msl.eval_country == c)]
    else:
        mod = msl[msl.model == f"{c}_adapter"]
    qids = sorted(popc.question_id.unique())
    by_dim = {d: [q for q in qids if qdim(q) == d] for d in DIMS}
    rows = []
    for d in DIMS:
        qs = by_dim[d]
        for k in range(0, max(len(qs), 1), 5):
            rows.append((d, qs[k:k + 5]))
    nrows = len(rows)
    fig = plt.figure(figsize=(16, 1.45 * nrows + 0.6))
    gs = fig.add_gridspec(nrows, 5, hspace=0.75, wspace=0.28,
                          left=0.03, right=0.99, top=0.94, bottom=0.04)
    shown_legend = False
    for i, (dim, qs) in enumerate(rows):
        fig.text(0.005, 1 - (i + 0.5) / nrows, dim, rotation=90, va="center",
                 ha="center", fontsize=9, fontweight="bold")
        for j, q in enumerate(qs):
            ax = fig.add_subplot(gs[i, j])
            pp = popc[(popc.question_id == q)].set_index("option_code")["population_prob"]
            mp = mod[(mod.question_id == q)].set_index("option_code")["model_prob"]
            codes = sorted(set(pp.index) | set(mp.index))
            ax.plot([float(x) for x in pp.index], pp.values, "o-", color=BLUE,
                    ms=3, lw=1.2, label="Target population" if not shown_legend else None)
            ax.plot([float(x) for x in mp.index], mp.values, "s-", color=ORANGE,
                    ms=3, lw=1.2, label=f"{c}_adapter" if not shown_legend else None)
            if not shown_legend:
                ax.legend(fontsize=7, loc="upper right", frameon=False)
                shown_legend = True
            ax.set_ylim(0, 1)
            ax.set_title(f"{q}  {qtext(q)}", fontsize=7, pad=2)
            ax.tick_params(labelsize=6)
            ax.set_xlabel("", fontsize=6)
            ax.spines[["top", "right"]].set_visible(False)
        for j in range(len(qs), 5):
            ax = fig.add_subplot(gs[i, j]); ax.axis("off")
    fig.suptitle(f"{c}_adapter vs Target population ({c}) — WVS wave 7 option distributions",
                 fontsize=12, fontweight="bold")
    fig.savefig(path, dpi=170)
    plt.close(fig)

plot_files = {}
for c in CCS:
    p = OUT / f"{c}_adapter_distributions.png"
    adapter_grid(c, p)
    plot_files[c] = p
    print("plotted", c)

# ---------------------------------------------------------------- key graphs
def graph_delta():
    fig, ax = plt.subplots(figsize=(9, 4.2))
    m = mt.sort_values("delta")
    colors = [GREEN if v >= 0 else RED for v in m.delta]
    ax.bar(m.index, m.delta, color=colors)
    ax.axhline(0, color="black", lw=0.8)
    ax.set_ylabel("Δ TVD (adapter − base); positive = adapter better")
    ax.set_title("Matched OOS alignment: adapter vs base model (TVD, 35 WVS items)")
    ax.bar_label(ax.containers[0], fmt="%+.3f", fontsize=8)
    fig.tight_layout(); p = OUT / "graph_delta_tvd.png"; fig.savefig(p, dpi=170); plt.close(fig)
    return p

def graph_heatmap():
    sub = imp[(imp.metric == "tv_distance") & (imp.relationship == "matched")]
    piv = sub.pivot_table(index="model", columns="gps_dimension",
                          values="mean_improvement_over_base", aggfunc="mean")
    piv = piv.reindex([f"{c}_adapter" for c in CCS]).reindex(columns=DIMS)
    fig, ax = plt.subplots(figsize=(10, 5))
    im = ax.imshow(piv.values, cmap="RdYlGn", vmin=-0.4, vmax=0.6, aspect="auto")
    ax.set_xticks(range(len(DIMS)), DIMS, rotation=30, ha="right")
    ax.set_yticks(range(len(CCS)), [c[:3] for c in CCS])
    for i in range(len(CCS)):
        for j in range(len(DIMS)):
            v = piv.values[i, j]
            if np.isfinite(v):
                ax.text(j, i, f"{v:+.2f}", ha="center", va="center", fontsize=7)
    ax.set_title("Δ TVD vs base by GPS dimension (matched; positive = adapter better)")
    fig.colorbar(im, shrink=0.8)
    fig.tight_layout(); p = OUT / "graph_dim_heatmap.png"; fig.savefig(p, dpi=170); plt.close(fig)
    return p

def graph_two_layers():
    fig, ax = plt.subplots(figsize=(8, 5))
    x = [RR_ACC[c] for c in CCS]; y = [mt.loc[c, "delta"] for c in CCS]
    ax.scatter(x, y, s=70)
    for c in CCS:
        ax.annotate(c, (RR_ACC[c], mt.loc[c, "delta"]), fontsize=8, xytext=(4, 4),
                    textcoords="offset points")
    ax.axhline(0, color="black", lw=0.8, ls="--")
    ax.set_xlabel("Reward recovery accuracy (held-out DPO pairs)")
    ax.set_ylabel("Δ TVD vs base (WVS OOS)")
    ax.set_title("Two layers: preference recovery vs population alignment")
    fig.tight_layout(); p = OUT / "graph_two_layers.png"; fig.savefig(p, dpi=170); plt.close(fig)
    return p

def graph_frac():
    fig, ax = plt.subplots(figsize=(9, 4.2))
    m = mt.sort_values("frac_improved")
    ax.bar(m.index, m.frac_improved, color="#7f7f7f")
    ax.set_ylim(0, 1)
    ax.set_ylabel("Fraction of questions where adapter < base TVD")
    ax.set_title("Per-question improvement rate (matched)")
    ax.bar_label(ax.containers[0], fmt="%.0f%%", fontsize=8)
    fig.tight_layout(); p = OUT / "graph_frac.png"; fig.savefig(p, dpi=170); plt.close(fig)
    return p

g_delta = graph_delta(); g_heat = graph_heatmap(); g_2l = graph_two_layers(); g_frac = graph_frac()

# ---------------------------------------------------------------- Ksennia-style doc
doc1 = Document()
doc1.add_heading("CO2 Adapters — WVS Wave 7 Option-Distribution Grids", 0)
doc1.add_paragraph("Target population (blue) vs adapter (orange) response distributions per question, "
                   "grouped by GPS dimension. Base model: Llama-3.1-8B-Instruct + QLoRA-DPO adapter per country. "
                   "Unconditioned prompting (anti-leakage): each adapter emits a single fixed distribution, "
                   "scored against each country's survey-weighted population.")
for c in CCS:
    doc1.add_heading(f"{c}_adapter on {c}  —  matched TVD {mt.loc[c, 'adapter_tvd']:.3f} "
                     f"(base {mt.loc[c, 'base_tvd']:.3f}, Δ {mt.loc[c, 'delta']:+.3f})", level=1)
    doc1.add_picture(str(plot_files[c]), width=Inches(6.6))
doc1.save(OUT / "CO2_WVS_adapter_distributions.docx")
print("doc1 saved")

# ---------------------------------------------------------------- executive report
doc = Document()
doc.add_heading("CO2 Country Adapters: WVS OOS Evaluation — Executive Report", 0)
doc.add_paragraph("Date: 2026-08-19 · 8 QLoRA-DPO adapters on Llama-3.1-8B-Instruct "
                  "(IND, IDN, NGA, EGY, TUR, NLD, BRA, GRC) · 35 unseen WVS wave-7 items "
                  "(30 GPS-mapped + 5 demographic) · option log-likelihood → softmax → "
                  "distributional metrics vs survey-weighted populations · unconditioned "
                  "prompts (single fixed distribution per model). EGY: 3 items skipped "
                  "(Q69–71 not asked in Egypt; n=32).")

doc.add_heading("1 · Headline: matched alignment vs base model", level=1)
t = doc.add_table(rows=1, cols=7)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for i, h in enumerate(["Country", "Base TVD", "Adapter TVD", "Δ TVD", "Frac. improved", "Base JSD", "Adapter JSD"]):
    hdr[i].text = h
for c in CCS:
    r = mt.loc[c]
    row = t.add_row().cells
    vals = [c, f"{r.base_tvd:.3f}", f"{r.adapter_tvd:.3f}", f"{r.delta:+.3f}",
            f"{r.frac_improved:.0%}", f"{r.base_jsd:.3f}", f"{r.adapter_jsd:.3f}"]
    for i, v in enumerate(vals):
        row[i].text = v
doc.add_paragraph("Pooled: base TVD %.3f vs adapter %.3f (%d of 8 improve). "
                  "Δ>0 = adapter closer to its country's population than the base model."
                  % (mt.base_tvd.mean(), mt.adapter_tvd.mean(), int((mt.delta > 0).sum())))
doc.add_picture(str(g_delta), width=Inches(6.4))

doc.add_heading("2 · Dimension structure", level=1)
doc.add_picture(str(g_heat), width=Inches(6.6))
doc.add_paragraph("Altruism is the strongest win dimension (BRA +0.446, GRC +0.588, both 100% of "
                  "questions improved); patience wins for IND/GRC/NGA/TUR. Negative reciprocity is the "
                  "consistent loser (NLD −0.363 with 0% improved, TUR −0.157, NGA −0.138). Trust is "
                  "mixed (NGA +0.079/83% vs IND −0.111). NLD degrades on every dimension.")

doc.add_heading("3 · Two layers: preference recovery vs population alignment", level=1)
doc.add_picture(str(g_2l), width=Inches(5.6))
doc.add_paragraph("Reward recovery (held-out DPO pairs, 88.6–97.7%) is a sanity layer: it measures "
                  "whether DPO moved the policy toward the country's synthetic labels. OOS TVD is the "
                  "validity layer: whether the resulting distribution resembles the real population. "
                  "The layers are weakly related — high recovery does not imply alignment (IDN/EGY "
                  "recover at 88.6% yet degrade on TVD), while NGA and GRC align well. Recovery "
                  "validates the plumbing; only calibration validates the science.")
doc.add_picture(str(g_frac), width=Inches(6.4))

doc.add_heading("4 · Robustness & caveats", level=1)
doc.add_paragraph("Bootstrap CIs (2,000 reps, per-question resampling) are in "
                  "survey_metric_bootstrap_summary.csv; ECE in survey_probability_calibration_ece.csv "
                  "(per-model). Duplicate-preference profiles (sign(z) labeling): IND≡GRC≡LTU and "
                  "IDN≡EGY have identical training labels — their adapters are effectively one model "
                  "(IDN_adapter on EGY == EGY_adapter on EGY = 0.578 exactly). EGY scored on 32 items. "
                  "Unconditioned prompts: matched-vs-cross compares each model's single fixed "
                  "distribution against each population (distributional alignment, not "
                  "country-conditioned inference).")

doc.add_heading("5 · Comparison with Ksennia's set (CHN, JPN, GBR, US, MEX, ARG, DEU, RUS)", level=1)
cmp = doc.add_table(rows=1, cols=4)
cmp.style = "Light Grid Accent 1"
for i, h in enumerate(["Dimension", "Ksennia's set (qualitative)", "CO2 set (quantitative)", "Consistency"]):
    cmp.rows[0].cells[i].text = h
comp_rows = [
    ("Trust", "Best dimension (MEX/ARG/DEU near-perfect); worst for JPN",
     "Mixed: NGA +0.079/83%, TUR +0.016; IND −0.111, IDN/EGY −0.085", "Dimension-selective, country-heterogeneous"),
    ("Altruism", "Worst dimension (inversions in CHN/USA/DEU/RUS)",
     "Best win dimension (BRA +0.446, GRC +0.588, 100%)", "Contradicts — CO2 adapters align where hers invert"),
    ("Negative reciprocity", "Poor (mismatched shapes)",
     "Consistent loser (NLD −0.363, TUR −0.157, NGA −0.138)", "Reinforced"),
    ("Patience", "Poor (inversions in GBR/USA)",
     "Win for 4/8 (IND +0.150, GRC +0.162); mild loss otherwise", "Mixed"),
    ("Risk taking", "Poor (GBR best); USA near-perfect",
     "Mild loss everywhere (−0.01 to −0.07)", "Reinforced (weak)"),
    ("Overall", "USA degrades; MEX selective (trust great, values poor)",
     "5/8 degrade (NLD −0.093 worst), 3/8 improve (GRC +0.075 best)",
     "Reinforced: degradation is the modal outcome; not universal"),
]
for dim, k, o, cons in comp_rows:
    row = cmp.add_row().cells
    for i, v in enumerate([dim, k, o, cons]):
        row[i].text = v
doc.add_paragraph("Notes: Ksennia's column is a qualitative reading of her distribution grids "
                  "(her quantitative CSVs were not in the shared document); CO2 numbers are paired "
                  "per-question mean TVD deltas vs base on matched questions. The USA-adapter "
                  "degradation from the preliminary 2×2 replicates at scale; the new insight is that "
                  "the failure is dimension- and country-specific rather than universal — GRC (and "
                  "BRA's altruism dimension, +0.446/100%) show the mechanism can work when the "
                  "synthetic sign signal aligns with the population's direction.")
doc.save(OUT / "CO2_WVS_executive_report.docx")
print("doc2 saved")
print("ALL DONE")
